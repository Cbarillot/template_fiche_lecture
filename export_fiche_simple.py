import json
import os
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from datetime import datetime
import sys

def creer_pdf(data, output_path):
    """Crée un fichier PDF à partir des données."""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # En-tête
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, 'Fiche de Lecture', 0, 1, 'C')
        
        # Sous-titre avec le titre de l'œuvre
        if data.get('titre'):
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, data['titre'], 0, 1, 'C')
        
        pdf.ln(10)
        
        # Police normale
        pdf.set_font('Arial', '', 12)
        
        # Contenu
        for cle, valeur in data.items():
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, f"{cle.capitalize()} :", ln=True)
            pdf.set_font('Arial', '', 12)
            
            if cle == 'citations':
                if valeur and any(cit.get('text') for cit in valeur):
                    for citation in valeur:
                        if citation.get('text'):
                            pdf.multi_cell(0, 8, f"- {citation['text']} (p.{citation.get('page', '?')})")
                        else:
                            pdf.multi_cell(0, 8, "- [Aucune citation renseignée]")
                else:
                    pdf.multi_cell(0, 8, "[Aucune citation renseignée]")
            elif isinstance(valeur, str):
                if valeur.strip():
                    pdf.multi_cell(0, 8, valeur)
                else:
                    pdf.multi_cell(0, 8, "[Non renseigné]")
            else:
                pdf.multi_cell(0, 8, str(valeur) if valeur else "[Non renseigné]")
                
            pdf.ln(5)
        
        # Enregistrement
        pdf.output(output_path)
        return True
    except Exception as e:
        print(f"Erreur lors de la création du PDF : {e}")
        return False

def creer_docx(data, output_path):
    """Crée un fichier DOCX à partir des données."""
    try:
        doc = Document()
        
        # Titre
        doc.add_heading('Fiche de Lecture', 0)
        
        # Sous-titre avec le titre de l'œuvre
        if data.get('titre'):
            doc.add_heading(data['titre'], level=1)
        
        doc.add_paragraph()
        
        # Contenu
        for cle, valeur in data.items():
            doc.add_heading(cle.capitalize(), level=2)
            
            if cle == 'citations':
                if valeur and any(cit.get('text') for cit in valeur):
                    for citation in valeur:
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        if citation.get('text'):
                            p.add_run(f"{citation['text']} ")
                            p.add_run(f"(p.{citation.get('page', '?')})").italic = True
                        else:
                            p.add_run("[Aucune citation renseignée]")
                else:
                    doc.add_paragraph("[Aucune citation renseignée]")
            elif isinstance(valeur, str):
                if valeur.strip():
                    doc.add_paragraph(valeur)
                else:
                    doc.add_paragraph("[Non renseigné]")
            else:
                doc.add_paragraph(str(valeur) if valeur else "[Non renseigné]")
        
        # Enregistrement
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Erreur lors de la création du DOCX : {e}")
        return False

def main():
    # Vérifier si un fichier a été fourni en argument
    if len(sys.argv) < 2:
        print("Utilisation : python export_fiche_simple.py chemin/vers/votre/fiche.json")
        return
    
    json_file = sys.argv[1]
    
    # Vérifier si le fichier existe
    if not os.path.exists(json_file):
        print(f"Erreur : Le fichier {json_file} n'existe pas.")
        return
    
    # Charger les données JSON
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier JSON : {e}")
        return
    
    # Créer le dossier d'export s'il n'existe pas
    export_dir = os.path.join(os.path.dirname(os.path.abspath(json_file)), 'exports')
    os.makedirs(export_dir, exist_ok=True)
    
    # Générer un nom de base pour les fichiers de sortie
    base_name = f"fiche_lecture_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
    # Exporter en PDF
    pdf_path = os.path.join(export_dir, f"{base_name}.pdf")
    if creer_pdf(data, pdf_path):
        print(f"✓ Fichier PDF créé : {pdf_path}")
    
    # Exporter en DOCX
    docx_path = os.path.join(export_dir, f"{base_name}.docx")
    if creer_docx(data, docx_path):
        print(f"✓ Fichier DOCX créé : {docx_path}")
    
    print("\nExportation terminée ! Les fichiers ont été enregistrés dans le dossier 'exports'.")

if __name__ == "__main__":
    main()
